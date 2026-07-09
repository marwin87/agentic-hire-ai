"""Tests for CV route handler functions."""

import hashlib
import pytest
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import HTTPException

from src.api.routes.cv import (
    _ingest_cv_background,
    calculate_file_hash,
    cv_status,
    upload_cv,
)

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _build_docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Real CV content.")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ===== calculate_file_hash =====


@pytest.mark.asyncio
async def test_calculate_file_hash_returns_sha256() -> None:
    content = b"hello world"
    result = await calculate_file_hash(content)
    expected = hashlib.sha256(b"hello world").hexdigest()
    assert result == expected


@pytest.mark.asyncio
async def test_calculate_file_hash_empty_bytes() -> None:
    result = await calculate_file_hash(b"")
    assert result == hashlib.sha256(b"").hexdigest()


# ===== cv_status =====


def _user(uid: object = None) -> MagicMock:
    u = MagicMock()
    u.id = uid or uuid4()
    return u


@pytest.mark.asyncio
async def test_cv_status_no_cv_returns_has_cv_false() -> None:
    with patch(
        "src.api.routes.cv.CVFileRepository.get_latest_by_user",
        new_callable=AsyncMock,
        return_value=None,
    ):
        result = await cv_status(user=_user(), db=AsyncMock())
    assert result["has_cv"] is False
    assert result["ingestion_status"] == "none"
    assert result["filename"] is None


@pytest.mark.asyncio
async def test_cv_status_completed_ingestion() -> None:
    cv_file = MagicMock()
    cv_file.file_path = "/data/cv/user123/resume_20260101.pdf"
    cv_file.ingested_at = datetime.now(timezone.utc)
    cv_file.ingestion_error = None

    with patch(
        "src.api.routes.cv.CVFileRepository.get_latest_by_user",
        new_callable=AsyncMock,
        return_value=cv_file,
    ):
        result = await cv_status(user=_user(), db=AsyncMock())

    assert result["has_cv"] is True
    assert result["ingestion_status"] == "completed"
    assert result["filename"] == "resume_20260101.pdf"


@pytest.mark.asyncio
async def test_cv_status_failed_ingestion() -> None:
    cv_file = MagicMock()
    cv_file.file_path = "/data/cv/user/resume.pdf"
    cv_file.ingested_at = None
    cv_file.ingestion_error = "Not a CV"

    with patch(
        "src.api.routes.cv.CVFileRepository.get_latest_by_user",
        new_callable=AsyncMock,
        return_value=cv_file,
    ):
        result = await cv_status(user=_user(), db=AsyncMock())

    assert result["ingestion_status"] == "failed"
    assert result["ingestion_error"] == "Not a CV"


@pytest.mark.asyncio
async def test_cv_status_processing() -> None:
    cv_file = MagicMock()
    cv_file.file_path = "/data/cv/user/resume.pdf"
    cv_file.ingested_at = None
    cv_file.ingestion_error = None

    with patch(
        "src.api.routes.cv.CVFileRepository.get_latest_by_user",
        new_callable=AsyncMock,
        return_value=cv_file,
    ):
        result = await cv_status(user=_user(), db=AsyncMock())

    assert result["ingestion_status"] == "processing"


# ===== _ingest_cv_background =====


def _make_session_factory(cv_file: object = None) -> MagicMock:
    mock_session = AsyncMock()
    mock_session.get = AsyncMock(return_value=cv_file)
    mock_session.commit = AsyncMock()
    mock_session.__aenter__ = AsyncMock(return_value=mock_session)
    mock_session.__aexit__ = AsyncMock(return_value=False)

    mock_factory = MagicMock(return_value=mock_session)
    return mock_factory


@pytest.mark.asyncio
async def test_ingest_cv_background_success_stamps_ingested_at() -> None:
    cv_file = MagicMock()
    cv_file.ingested_at = None
    mock_factory = _make_session_factory(cv_file)

    mock_vm = AsyncMock()
    mock_vm.ingest_cv_async = AsyncMock()

    with patch("src.api.routes.cv.get_session_factory", return_value=mock_factory):
        await _ingest_cv_background(uuid4(), "path/cv.pdf", mock_vm)

    assert cv_file.ingested_at is not None


@pytest.mark.asyncio
async def test_ingest_cv_background_value_error_sets_ingestion_error() -> None:
    cv_file = MagicMock()
    cv_file.ingestion_error = None
    mock_factory = _make_session_factory(cv_file)

    mock_vm = AsyncMock()
    mock_vm.ingest_cv_async = AsyncMock(side_effect=ValueError("Not a CV document"))

    with patch("src.api.routes.cv.get_session_factory", return_value=mock_factory):
        await _ingest_cv_background(uuid4(), "path/cv.pdf", mock_vm)

    assert cv_file.ingestion_error == "Not a CV document"


@pytest.mark.asyncio
async def test_ingest_cv_background_generic_exception_sets_generic_error() -> None:
    cv_file = MagicMock()
    cv_file.ingestion_error = None
    mock_factory = _make_session_factory(cv_file)

    mock_vm = AsyncMock()
    mock_vm.ingest_cv_async = AsyncMock(side_effect=RuntimeError("GPU crash"))

    with patch("src.api.routes.cv.get_session_factory", return_value=mock_factory):
        await _ingest_cv_background(uuid4(), "path/cv.pdf", mock_vm)

    assert "failed" in cv_file.ingestion_error.lower()


# ===== upload_cv =====


def _upload_file(
    content_type: str = "application/pdf", content: bytes = b"PDF"
) -> MagicMock:
    f = AsyncMock()
    f.content_type = content_type
    f.read = AsyncMock(return_value=content)
    f.filename = "resume.pdf"
    return f


@pytest.mark.asyncio
async def test_upload_cv_invalid_content_type_raises_400() -> None:
    with pytest.raises(HTTPException) as exc:
        await upload_cv(
            background_tasks=MagicMock(),
            file=_upload_file(content_type="text/plain"),
            user=_user(),
            db=AsyncMock(),
        )
    assert exc.value.status_code == 400
    assert "pdf" in exc.value.detail.lower()


@pytest.mark.asyncio
async def test_upload_cv_empty_file_raises_400() -> None:
    with pytest.raises(HTTPException) as exc:
        await upload_cv(
            background_tasks=MagicMock(),
            file=_upload_file(content=b""),
            user=_user(),
            db=AsyncMock(),
        )
    assert exc.value.status_code == 400
    assert "empty" in exc.value.detail.lower()


@pytest.mark.asyncio
async def test_upload_cv_oversized_file_raises_400() -> None:
    big_content = b"%PDF" + b"X" * (10 * 1024 * 1024 + 1)
    with pytest.raises(HTTPException) as exc:
        await upload_cv(
            background_tasks=MagicMock(),
            file=_upload_file(content=big_content),
            user=_user(),
            db=AsyncMock(),
        )
    assert exc.value.status_code == 400
    assert "large" in exc.value.detail.lower() or "mb" in exc.value.detail.lower()


@pytest.mark.asyncio
async def test_upload_cv_success_queues_background_task(tmp_path: Path) -> None:
    content = b"%PDF-1.4 test content"
    user_id = uuid4()

    db = AsyncMock()
    db.add = MagicMock()
    db.flush = AsyncMock()
    db.commit = AsyncMock()
    db.delete = AsyncMock()
    background = MagicMock()

    with (
        patch.object(Path, "mkdir"),
        patch.object(Path, "write_bytes"),
        patch(
            "src.api.routes.cv.CVFileRepository.get_latest_by_user",
            new_callable=AsyncMock,
            return_value=None,
        ),
        patch("src.api.routes.cv.ChatOpenAI"),
        patch("src.api.routes.cv.OpenAIEmbeddings"),
        patch("src.api.routes.cv.CVVectorManager"),
    ):
        result = await upload_cv(
            background_tasks=background,
            file=_upload_file(content=content),
            user=_user(uid=user_id),
            db=db,
        )

    assert result["status"] == "processing"
    assert "file_hash" in result
    background.add_task.assert_called_once()


@pytest.mark.asyncio
async def test_upload_cv_valid_docx_queues_background_task(tmp_path: Path) -> None:
    content = _build_docx_bytes()
    user_id = uuid4()

    db = AsyncMock()
    db.add = MagicMock()
    db.flush = AsyncMock()
    db.commit = AsyncMock()
    db.delete = AsyncMock()
    background = MagicMock()

    with (
        patch.object(Path, "mkdir"),
        patch.object(Path, "write_bytes"),
        patch(
            "src.api.routes.cv.CVFileRepository.get_latest_by_user",
            new_callable=AsyncMock,
            return_value=None,
        ),
        patch("src.api.routes.cv.ChatOpenAI"),
        patch("src.api.routes.cv.OpenAIEmbeddings"),
        patch("src.api.routes.cv.CVVectorManager"),
    ):
        result = await upload_cv(
            background_tasks=background,
            file=_upload_file(content_type=DOCX_CONTENT_TYPE, content=content),
            user=_user(uid=user_id),
            db=db,
        )

    assert result["status"] == "processing"
    background.add_task.assert_called_once()
    saved_filepath = background.add_task.call_args.args[2]
    assert saved_filepath.endswith(".docx")


@pytest.mark.asyncio
async def test_upload_cv_docx_corrupt_internal_structure_raises_400() -> None:
    corrupt_content = b"PK\x03\x04" + b"not a real docx internal structure"
    with pytest.raises(HTTPException) as exc:
        await upload_cv(
            background_tasks=MagicMock(),
            file=_upload_file(content_type=DOCX_CONTENT_TYPE, content=corrupt_content),
            user=_user(),
            db=AsyncMock(),
        )
    assert exc.value.status_code == 400
    assert "invalid" in exc.value.detail.lower()


@pytest.mark.asyncio
async def test_upload_cv_docx_missing_zip_signature_raises_400() -> None:
    with pytest.raises(HTTPException) as exc:
        await upload_cv(
            background_tasks=MagicMock(),
            file=_upload_file(content_type=DOCX_CONTENT_TYPE, content=b"not a zip"),
            user=_user(),
            db=AsyncMock(),
        )
    assert exc.value.status_code == 400


# ===== Review fix regression tests =====
# F1: MAX_FILE_SIZE must be checked before any type-specific parsing (PDF
# magic bytes or DOCX unzip+XML parse), so an oversized upload is rejected
# on size alone even if its content would otherwise fail docx validation.
@pytest.mark.asyncio
async def test_upload_cv_oversized_docx_rejected_on_size_before_parsing() -> None:
    """Oversized DOCX-typed upload is rejected with the size message, not the
    'invalid file' message — proving the size gate runs before DocxDocument()
    ever attempts to parse the (here: garbage, non-docx) content.
    """
    oversized_garbage = b"not a valid docx at all" * (10 * 1024 * 1024)
    with (
        patch("src.api.routes.cv.DocxDocument") as mock_docx_ctor,
        pytest.raises(HTTPException) as exc,
    ):
        await upload_cv(
            background_tasks=MagicMock(),
            file=_upload_file(
                content_type=DOCX_CONTENT_TYPE, content=oversized_garbage
            ),
            user=_user(),
            db=AsyncMock(),
        )
    assert exc.value.status_code == 400
    assert "large" in exc.value.detail.lower()
    # The size gate must short-circuit before any DOCX parsing is attempted.
    mock_docx_ctor.assert_not_called()


# F2: the DOCX validation except clause is narrowed to
# (ValueError, PackageNotFoundError, BadZipFile). An unrelated exception type
# (simulating an internal bug, e.g. a future python-docx API change) must
# propagate rather than being silently swallowed as a generic 400.
@pytest.mark.asyncio
async def test_upload_cv_docx_unexpected_exception_propagates_not_swallowed() -> None:
    content = b"PK\x03\x04" + b"payload"
    with (
        patch(
            "src.api.routes.cv.DocxDocument",
            side_effect=RuntimeError("unexpected internal docx-library bug"),
        ),
        pytest.raises(RuntimeError, match="unexpected internal docx-library bug"),
    ):
        await upload_cv(
            background_tasks=MagicMock(),
            file=_upload_file(content_type=DOCX_CONTENT_TYPE, content=content),
            user=_user(),
            db=AsyncMock(),
        )
