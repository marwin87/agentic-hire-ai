"""Tests for CVVectorManager DOCX extraction and detection."""

import pytest
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

from docx import Document as DocxDocument

from src.tools.vectordb import CVVectorManager, CVDetectionResult, MIN_DOCX_TEXT_LENGTH


@pytest.fixture
def manager(tmp_path: Path) -> CVVectorManager:
    mock_vision = MagicMock()
    mock_embeddings = MagicMock()
    return CVVectorManager(
        vision_model=mock_vision,
        embeddings=mock_embeddings,
        user_id=uuid4(),
        cv_cache_dir=str(tmp_path),
    )


def _build_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("John Doe", level=0)  # Title
    doc.add_heading("Experience", level=1)  # Heading 1
    doc.add_heading("Acme Corp", level=2)  # Heading 2
    doc.add_heading("Details", level=3)  # Heading 3
    doc.add_paragraph("Built things and shipped code.")
    doc.add_paragraph("Wrote Python", style="List Bullet")
    doc.add_paragraph("   ")  # whitespace-only, should be skipped
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Skill"
    table.rows[0].cells[1].text = "Level"
    table.rows[1].cells[0].text = "Python"
    table.rows[1].cells[1].text = "Expert"
    doc.save(str(path))


# ===== _docx_to_markdown =====


def test_docx_to_markdown_maps_headings_and_bullets(tmp_path: Path) -> None:
    docx_path = tmp_path / "cv.docx"
    _build_docx(docx_path)

    result = CVVectorManager._docx_to_markdown(str(docx_path))

    assert "# John Doe" in result
    assert "# Experience" in result
    assert "## Acme Corp" in result
    assert "Details" in result  # Heading 3 too deep for our mapping, still present
    assert "Built things and shipped code." in result
    assert "• Wrote Python" in result


def test_docx_to_markdown_flattens_tables_to_pipe_separated_lines(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "cv.docx"
    _build_docx(docx_path)

    result = CVVectorManager._docx_to_markdown(str(docx_path))

    assert "Skill | Level" in result
    assert "Python | Expert" in result


def test_docx_to_markdown_skips_empty_paragraphs(tmp_path: Path) -> None:
    docx_path = tmp_path / "cv.docx"
    doc = DocxDocument()
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("Real content.")
    doc.save(str(docx_path))

    result = CVVectorManager._docx_to_markdown(str(docx_path))

    assert result.strip() == "Real content."


# ===== _detect_if_cv_from_text =====


def test_detect_if_cv_from_text_passes_when_is_cv_true(
    manager: CVVectorManager,
) -> None:
    mock_detector = MagicMock()
    mock_detector.invoke.return_value = CVDetectionResult(
        is_cv=True, reason="Looks like a resume."
    )
    manager.vision_model.with_structured_output.return_value = mock_detector

    manager._detect_if_cv_from_text("John Doe\nExperience\nAcme Corp")

    manager.vision_model.with_structured_output.assert_called_once_with(
        CVDetectionResult
    )


def test_detect_if_cv_from_text_raises_when_is_cv_false(
    manager: CVVectorManager,
) -> None:
    mock_detector = MagicMock()
    mock_detector.invoke.return_value = CVDetectionResult(
        is_cv=False, reason="This is a recipe, not a CV."
    )
    manager.vision_model.with_structured_output.return_value = mock_detector

    with pytest.raises(ValueError, match="recipe"):
        manager._detect_if_cv_from_text("Ingredients: flour, sugar, eggs.")


# ===== ingest_cv_async dispatch =====


def _mock_repo_and_session() -> AsyncMock:
    session = AsyncMock()
    session.__aenter__ = AsyncMock(return_value=session)
    session.__aexit__ = AsyncMock(return_value=False)
    factory = MagicMock(return_value=session)
    return factory


@pytest.mark.asyncio
async def test_ingest_cv_async_docx_dispatch_uses_docx_path(
    manager: CVVectorManager, tmp_path: Path
) -> None:
    docx_path = tmp_path / "cv.docx"
    _build_docx(docx_path)

    mock_detector = MagicMock()
    mock_detector.invoke.return_value = CVDetectionResult(is_cv=True, reason="CV")
    manager.vision_model.with_structured_output.return_value = mock_detector
    manager.embeddings.embed_query.return_value = [0.0]

    factory = _mock_repo_and_session()
    with (
        patch("src.tools.vectordb.get_session_factory", return_value=factory),
        patch(
            "src.tools.vectordb.CVEmbeddingRepository.delete_by_user",
            new_callable=AsyncMock,
        ),
        patch(
            "src.tools.vectordb.CVEmbeddingRepository.bulk_insert",
            new_callable=AsyncMock,
        ),
        patch.object(manager, "_pdf_to_base64_images") as mock_pdf_images,
        patch.object(manager, "_detect_if_cv") as mock_detect_if_cv,
    ):
        await manager.ingest_cv_async(str(docx_path))

    mock_pdf_images.assert_not_called()
    mock_detect_if_cv.assert_not_called()


@pytest.mark.asyncio
async def test_ingest_cv_async_docx_guard_rejects_sparse_text(
    manager: CVVectorManager, tmp_path: Path
) -> None:
    docx_path = tmp_path / "sparse.docx"
    doc = DocxDocument()
    doc.add_paragraph("short")
    doc.save(str(docx_path))

    assert len("short") < MIN_DOCX_TEXT_LENGTH

    with pytest.raises(ValueError, match="no extractable text"):
        await manager.ingest_cv_async(str(docx_path))


@pytest.mark.asyncio
async def test_ingest_cv_async_docx_package_not_found_raises_value_error(
    manager: CVVectorManager, tmp_path: Path
) -> None:
    bad_path = tmp_path / "corrupt.docx"
    bad_path.write_bytes(b"not a real docx file")

    with pytest.raises(ValueError, match="invalid or corrupted"):
        await manager.ingest_cv_async(str(bad_path))
